from docx import Document
from docx.shared import Pt

documentAnswerSheet = Document()
documentQuestionSheet = Document()
language1 = []
language2 = []

print('What should be the file name:')
fileName = input()

print('How many Vocabularies will there be:')
numVocab = input()

for x in range(int(numVocab)):
    print('Vocab in Language 1:')
    language1.append(input())
    print('Vocab in Language 2:')
    language2.append(input())

for x in range(int(numVocab)):
    run1 = documentAnswerSheet.add_paragraph().add_run(language1[x] + ' = ' + language2[x])
    run2 = documentQuestionSheet.add_paragraph().add_run(language1[x] + ' = ')
    run1.font.size = Pt(24)
    run2.font.size = Pt(24)


documentQuestionSheet.save(fileName + ' - Question Sheet.docx')
documentAnswerSheet.save(fileName + ' - Answer Sheet.docx')